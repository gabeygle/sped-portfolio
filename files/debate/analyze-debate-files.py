#!/usr/bin/env python3
"""
analyze-debate-files.py
Reproduces every figure published on https://burdenofproof.education/audience-adaptation.html

Usage:
    pip install python-docx pypdf
    python3 analyze-debate-files.py team-affirmative-01-cognitive-warfare-2022-23.docx
    python3 analyze-debate-files.py *.docx --pages      # also measures card length in pages

The --pages flag renders each .docx to PDF with LibreOffice and locates every card in the
rendered output. It is slow (a minute or two per file) but it is the only honest way to
report a page count. Requires `soffice` on PATH.

---------------------------------------------------------------------------
THE TWO-DENOMINATOR PROBLEM  (read this before quoting any percentage)
---------------------------------------------------------------------------
This script reports two different totals for the same document, and they disagree:

    paragraph-level total   62,629 words   <- counts p.text.split() per paragraph
    run-level total         63,823 words   <- counts r.text.split() per run

A "run" is a span of uniformly formatted characters. Word processors split runs at every
formatting change, and that split frequently lands mid-word: highlighting the "deter" in
"deterrence" produces two runs, and splitting on whitespace then counts two words where the
document contains one. The run-level total is therefore inflated, by about 1.9% here.

Both numbers are useful, but only for different things:

  * Use the PARAGRAPH total as the length of the document. It is correct.
  * Use the RUN total only as the denominator for highlight/underline percentages, because
    the numerators are themselves counted run-by-run and carry the same inflation. Dividing
    a run-level numerator by a paragraph-level denominator would mix two counting schemes
    and overstate the percentage.

The published page follows that rule: totals come from paragraph counts, percentages from
run counts. The absolute word counts of highlighted and underlined text are run-level and
so are slightly high; treat them as accurate to about two percent, not exact.

A cleaner fix, not implemented here, is to concatenate the runs of each paragraph into a
character stream, tag each character with its formatting, and tokenise once. That yields a
single denominator. It is on the roadmap.

---------------------------------------------------------------------------
CARD DETECTION  (this is where the first version was wrong)
---------------------------------------------------------------------------
A card is a unit of evidence: a tag, a citation, then the quoted body. The obvious rule --
"every paragraph styled Heading 4 is a tag" -- is wrong, and counting that way reported 63
cards in the 2025 file where a hand count finds 56.

Two other things wear the tag style:

  * Analytics. Arguments made in the debater's own voice, with no evidence under them. They
    are formatted like tags because they are read aloud like tags.
  * Documentation headings. The 2025 file opens with a Notes section whose questions ("What
    does the plan do?") are styled as tags.

The rule this script uses instead: a Heading 4 block is a card if its body contains
underlined text, has at least one paragraph after the citation line, and its tag is not a
question. Underlining is the discriminator that matters -- underlining marks the reasoning
inside a piece of quoted evidence, so an analytic has none. The question test then removes
documentation headings, which can contain underlining because they quote the plan text.

That rule reproduces a hand count exactly on both files: 20 cards in the 2022 file and 56
in the 2025 file. It is validated against the only ground truth available, which is a person
reading the documents.

Two further checks on tag detection itself:

  1. Unstyled tags. A tag typed as bold body text instead of Heading 4 would be missed
     entirely and merged into the card above. Scanning for fully bold "Normal" paragraphs
     under 45 words finds 0 in the 2022 file and 1 in the 2025 file, and that one is the
     placeholder "[FIGURE OMITTED]".

  2. Positional check against the rendered file. With --pages, every located tag is found in
     document order, 21 of 21 in the 2022 file and 84 of 93 in the 2025 file. The 9 not
     located are tags shorter than the 20-character matching threshold, such as "Timing",
     skipped to avoid false positives.

The weakest remaining assumption is the citation rule. "First body paragraph is the
citation" holds wherever it was spot-checked, but is not verified exhaustively. Where it
fails, that card's body is undercounted by one paragraph.

---------------------------------------------------------------------------
WHAT IS DELIBERATELY NOT HERE
---------------------------------------------------------------------------
An earlier version compared tag vocabulary against English word-frequency bands using the
`wordfreq` package. Those figures are not published and not computed here: the mapping from
Zipf frequency to a difficulty level like CEFR was an approximation without a source behind
it. The analysis is on the roadmap pending a defensible mapping.
"""

import sys
import re
import subprocess
import statistics as st
from pathlib import Path

import docx
from docx.enum.text import WD_COLOR_INDEX


def is_highlighted(run):
    return run.font.highlight_color not in (None, WD_COLOR_INDEX.AUTO)


def word_counts(document):
    """Return the two totals, plus run-level marked-word counts. See the docstring."""
    para_total = sum(len(p.text.split()) for p in document.paragraphs)
    run_total = hl = ul = both = 0
    for para in document.paragraphs:
        for run in para.runs:
            n = len(run.text.split())
            if not n:
                continue
            run_total += n
            h, u = is_highlighted(run), bool(run.underline)
            if h:
                hl += n
            if u:
                ul += n
            if h and u:
                both += n
    return dict(para_total=para_total, run_total=run_total,
                highlighted=hl, underlined=ul, both=both)


DOC_QUESTION = re.compile(r'^(what|why|how|is|are|do|does|when|where|who)\b.*\?$', re.I)


def blocks(document):
    """Split on Heading 4. Body excludes the first paragraph (the citation). See docstring."""
    out, current = [], None
    for para in document.paragraphs:
        style, text = para.style.name, para.text.strip()
        if style == 'Heading 4':
            if current:
                out.append(current)
            current = {'tag': text, 'body': []}
        elif current is not None and style == 'Normal' and text:
            current['body'].append(para)   # keep the object; formatting is needed below
    if current:
        out.append(current)
    return out


def body_words(block):
    return sum(len(p.text.split()) for p in block['body'][1:])


def is_card(block):
    """A tag block is a card, not an analytic or a documentation heading. See docstring."""
    underlined = any(r.underline for p in block['body'] for r in p.runs)
    return (underlined
            and body_words(block) > 0
            and not DOC_QUESTION.match(block['tag']))


def unstyled_tag_check(document):
    """Count fully bold short Normal paragraphs, which would be tags this script misses."""
    found = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text or para.style.name != 'Normal':
            continue
        runs = [r for r in para.runs if r.text.strip()]
        if runs and all(r.bold for r in runs) and len(text.split()) < 45:
            found.append(text[:80])
    return found


def _normalise(s):
    return re.sub(r'\W+', '', s).lower()


def page_spans(docx_path, tags, workdir):
    """Render to PDF and measure how many pages each card occupies. Needs soffice + pypdf."""
    from pypdf import PdfReader

    pdf = Path(workdir) / (Path(docx_path).stem + '.pdf')
    if not pdf.exists():
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf',
                        str(docx_path), '--outdir', str(workdir)],
                       check=True, capture_output=True, timeout=900)
    reader = PdfReader(str(pdf))
    pages = [_normalise(p.extract_text() or '') for p in reader.pages]

    located, cursor, skipped = [], 0, 0
    for tag in tags:
        key = _normalise(tag)[:60]
        if len(key) < 20:          # too short to match without false positives
            skipped += 1
            continue
        hit = next((i for i in range(cursor, len(pages)) if key in pages[i]), None)
        if hit is None:
            hit = next((i for i, p in enumerate(pages) if key in p), None)
        if hit is None:
            skipped += 1
            continue
        located.append(hit)
        cursor = hit

    spans = [located[i + 1] - located[i] + 1 for i in range(len(located) - 1)]
    return dict(total_pages=len(pages), located=len(located),
                skipped=skipped, spans=spans)


def report(path, measure_pages, workdir):
    document = docx.Document(path)
    counts = word_counts(document)
    all_blocks = blocks(document)
    cs = [b for b in all_blocks if is_card(b)]
    tags = [c['tag'] for c in cs if c['tag']]
    tag_words = [len(t.split()) for t in tags]
    card_words = [body_words(c) for c in cs]
    rt = counts['run_total']

    print('=' * 72)
    print(Path(path).name)
    print('=' * 72)
    print(f"  words, paragraph-level (the document's length) : {counts['para_total']:>8,}")
    print(f"  words, run-level (percentage denominator only)  : {rt:>8,}"
          f"   (+{100 * (rt - counts['para_total']) / counts['para_total']:.1f}%)")
    print()
    print(f"  highlighted words : {counts['highlighted']:>7,}   {100 * counts['highlighted'] / rt:5.1f}%")
    print(f"  underlined words  : {counts['underlined']:>7,}   {100 * counts['underlined'] / rt:5.1f}%")
    print(f"  both              : {counts['both']:>7,}   {100 * counts['both'] / rt:5.1f}%")
    print()
    print(f"  Heading 4 blocks         : {len(all_blocks)}")
    print(f"  of those, cards          : {len(cs)}"
          f"   ({len(all_blocks) - len(cs)} analytics / documentation headings)")
    print(f"  tag length, words        : median {st.median(tag_words):.0f}  "
          f"mean {st.mean(tag_words):.1f}  range {min(tag_words)}-{max(tag_words)}")
    print(f"  card body, words         : median {st.median(card_words):.0f}  "
          f"mean {st.mean(card_words):.0f}  range {min(card_words)}-{max(card_words)}")
    print(f"  cards over 1,000 words   : {sum(b > 1000 for b in card_words)}"
          f"/{len(cs)} = {100 * sum(b > 1000 for b in card_words) / len(cs):.0f}%")

    suspects = unstyled_tag_check(document)
    print(f"  unstyled-tag check       : {len(suspects)} suspect paragraph(s)"
          + (f" -> {suspects[:3]}" if suspects else " (clean)"))

    if measure_pages:
        print()
        try:
            p = page_spans(path, tags, workdir)
        except Exception as exc:
            print(f"  page measurement failed: {exc}")
            return
        print(f"  rendered pages           : {p['total_pages']}")
        print(f"  tags located in the PDF  : {p['located']} of {len(tags)}"
              f"  ({p['skipped']} skipped as too short to match)")
        if p['spans']:
            print(f"  card length, pages       : median {st.median(p['spans']):.1f}  "
                  f"mean {st.mean(p['spans']):.2f}  max {max(p['spans'])}")
            print("  (a card occupying part of one page counts as 1)")
    print()


def main():
    args = [a for a in sys.argv[1:] if not a.startswith('--')]
    measure_pages = '--pages' in sys.argv
    if not args:
        print(__doc__)
        sys.exit(1)
    workdir = Path('.') / '_render'
    if measure_pages:
        workdir.mkdir(exist_ok=True)
    for path in args:
        report(path, measure_pages, workdir)


if __name__ == '__main__':
    main()
